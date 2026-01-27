# app/extract.py
"""Document text extraction with page tracking and quality validation."""

import json
import re
from dataclasses import dataclass
from datetime import datetime
from datetime import timezone
from pathlib import Path

import fitz
from docx import Document

from app.logging import get_logger

log = get_logger(__name__)


class ExtractionError(Exception):
    """Raised when document extraction fails."""

    pass


@dataclass
class PageContent:
    """Content extracted from a single page."""

    page_num: int
    text: str


@dataclass
class ExtractedDocument:
    """Extracted document with pages and metadata."""

    pages: list[PageContent]
    page_count: int
    source_file: str
    extraction_method: str

    @property
    def full_text(self) -> str:
        """Get the full document text."""
        return "\n".join(page.text for page in self.pages)

    @property
    def word_count(self) -> int:
        """Get total word count."""
        return sum(len(page.text.split()) for page in self.pages)

    @property
    def is_empty(self) -> bool:
        """Check if document has no extractable content."""
        return self.word_count < 10


def extract_pdf(path: Path) -> ExtractedDocument:
    """Extract text and metadata from a PDF file.

    Args:
        path: Path to the PDF file.

    Returns:
        ExtractedDocument with pages and metadata.

    Raises:
        FileNotFoundError: If the file does not exist.
        ExtractionError: If PDF extraction fails or document is corrupted.
    """
    if not path.exists():
        raise FileNotFoundError(f"PDF file not found: {path}")

    try:
        doc = fitz.open(path)
    except Exception as e:
        log.error("pdf_open_failed", filename=path.name, error=str(e))
        raise ExtractionError(
            f"Cannot open PDF {path.name}: corrupted or invalid"
        ) from e

    try:
        pages = []
        for page in doc:
            pages.append(
                PageContent(
                    page_num=page.number + 1,  # 1-indexed
                    text=page.get_text(),
                )
            )
        doc.close()

        extracted = ExtractedDocument(
            pages=pages,
            page_count=len(pages),
            source_file=path.name,
            extraction_method="pymupdf",
        )

        # Validate extraction quality
        if extracted.is_empty:
            log.warning(
                "extraction_empty",
                filename=path.name,
                pages=len(pages),
                words=extracted.word_count,
            )

        log.debug(
            "pdf_extracted",
            filename=path.name,
            pages=extracted.page_count,
            words=extracted.word_count,
        )
        return extracted

    except Exception as e:
        doc.close()
        log.error("pdf_extraction_failed", filename=path.name, error=str(e))
        raise ExtractionError(f"Failed to extract PDF {path.name}: {e}") from e


def extract_docx(path: Path) -> ExtractedDocument:
    """Extract text from a DOCX file.

    Note: DOCX files don't have page numbers in the same way as PDFs.
    Each paragraph is treated as belonging to page 1.

    Args:
        path: Path to the DOCX file.

    Returns:
        ExtractedDocument with content and metadata.

    Raises:
        FileNotFoundError: If the file does not exist.
        ExtractionError: If DOCX extraction fails.
    """
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")

    try:
        doc = Document(str(path))

        # Extract paragraphs
        content_parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                content_parts.append(p.text)

        # Extract tables (often contain fee schedules, definitions, etc.)
        for table in doc.tables:
            table_rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                # Dedupe adjacent cells (merged cells repeat)
                deduped: list[str] = []
                for c in cells:
                    if not deduped or c != deduped[-1]:
                        deduped.append(c)
                if any(deduped):  # Skip empty rows
                    table_rows.append(" | ".join(deduped))
            if table_rows:
                content_parts.append("\n".join(table_rows))

        full_text = "\n".join(content_parts)

        # DOCX doesn't have page boundaries, treat as single page
        pages = [PageContent(page_num=1, text=full_text)]

        extracted = ExtractedDocument(
            pages=pages,
            page_count=1,
            source_file=path.name,
            extraction_method="python-docx",
        )

        if extracted.is_empty:
            log.warning(
                "extraction_empty", filename=path.name, words=extracted.word_count
            )

        log.debug("docx_extracted", filename=path.name, words=extracted.word_count)
        return extracted
    except Exception as e:
        log.error("docx_extraction_failed", filename=path.name, error=str(e))
        raise ExtractionError(f"Failed to extract DOCX {path.name}: {e}") from e


def extract_document(path: Path) -> ExtractedDocument:
    """Extract text from a document based on its extension.

    Args:
        path: Path to the document file.

    Returns:
        ExtractedDocument with content and metadata.

    Raises:
        ExtractionError: If extraction fails or file type is not supported.
    """
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    elif suffix == ".docx":
        return extract_docx(path)
    else:
        log.error("unsupported_file_type", filename=path.name, suffix=suffix)
        raise ExtractionError(f"Unsupported file type: {suffix}")


def validate_extraction(extracted: ExtractedDocument) -> list[str]:
    """Validate extraction quality and return warnings.

    Args:
        extracted: The extracted document to validate.

    Returns:
        List of warning messages. Empty list if no issues.
    """
    warnings = []

    if extracted.is_empty:
        warnings.append(f"Document '{extracted.source_file}' has no extractable text")

    if extracted.page_count == 0:
        warnings.append(f"Document '{extracted.source_file}' has no pages")

    # Check for potential extraction issues (mostly non-alphabetic characters)
    # This can indicate: scanned PDF needing OCR, or font encoding issues
    text = extracted.full_text
    if text:
        alpha_ratio = sum(1 for c in text if c.isalpha()) / max(len(text), 1)
        if alpha_ratio < 0.3:
            warnings.append(
                f"Document '{extracted.source_file}' has extraction issues "
                f"(only {alpha_ratio:.0%} alphabetic) - may need OCR or has font encoding problems"
            )

    for warning in warnings:
        log.warning("extraction_validation", message=warning)

    return warnings


def detect_document_version(text: str) -> str | None:
    """Attempt to detect document version from text content.

    Looks for common version patterns like:
    - Version 1.0, Version 2.3.4
    - v1.0, v2.3
    - Effective Date patterns

    Args:
        text: The document text to search.

    Returns:
        Detected version string or None if not found.
    """

    # Check first 2000 chars (usually contains version info)
    sample = text[:2000]

    # Common version patterns (case-insensitive)
    patterns = [
        r"[Vv][Ee][Rr][Ss][Ii][Oo][Nn]\s*:?\s*(\d+(?:\.\d+)*)",  # Version 1.0, VERSION: 2.3.4
        r"\b[Vv](\d+(?:\.\d+)+)\b",  # v1.0, v2.3
        r"[Rr][Ee][Vv][Ii][Ss][Ii][Oo][Nn]\s*:?\s*(\d+(?:\.\d+)*)",  # Revision 1.0
    ]

    for pattern in patterns:
        match = re.search(pattern, sample)
        if match:
            return match.group(1)

    return None


def save_extraction_artifacts(
    extracted: ExtractedDocument,
    output_dir: Path,
    provider: str,
    relative_path: Path | None = None,
) -> tuple[Path, Path]:
    """Save extraction artifacts (.txt and .meta.json) as per spec.

    Args:
        extracted: The extracted document.
        output_dir: Directory to save artifacts (e.g., data/text/cme/).
        provider: Provider identifier (e.g., "cme").
        relative_path: Relative path from provider raw directory (for subdirectory support).

    Returns:
        Tuple of (text_path, meta_path) for saved files.
    """

    output_dir.mkdir(parents=True, exist_ok=True)

    # Use relative path if provided (encoded with __), else fall back to basename
    # e.g., "Fees/document.pdf" -> "Fees__document.pdf.txt"
    # This prevents collisions when same filename exists in different subdirectories
    if relative_path:
        source_name = str(relative_path).replace("/", "__")
    else:
        source_name = Path(extracted.source_file).name

    # Save text file
    text_path = output_dir / f"{source_name}.txt"
    text_path.write_text(extracted.full_text, encoding="utf-8")

    # Save metadata JSON
    meta_path = output_dir / f"{source_name}.meta.json"
    metadata = {
        "source_file": extracted.source_file,
        "provider": provider,
        "extracted_at": datetime.now(timezone.utc).isoformat(),
        "page_count": extracted.page_count,
        "extraction_method": extracted.extraction_method,
        "word_count": extracted.word_count,
        "document_version": detect_document_version(extracted.full_text),
    }
    meta_path.write_text(json.dumps(metadata, indent=2), encoding="utf-8")

    log.debug(
        "extraction_artifacts_saved",
        text_file=str(text_path),
        meta_file=str(meta_path),
    )

    return text_path, meta_path
